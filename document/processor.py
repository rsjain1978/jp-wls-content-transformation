"""Document processing functionality"""

import os
import asyncio
import base64
from io import BytesIO
from typing import List, Tuple
from termcolor import colored
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from openai import AsyncOpenAI
from pathlib import Path
import json
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process documents using GPT-4o for text extraction and translation"""
    
    def __init__(self):
        self.client = AsyncOpenAI()
        self.temp_dir = Path("static/temp")
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
    async def process_pdf(self, file_path: str, target_language: str) -> str:
        """Process PDF document and translate content"""
        try:
            logger.info(f"Processing PDF: {file_path}")
            
            # Convert PDF to images and store paths
            images, image_paths = await self._pdf_to_images(file_path)
            logger.info(f"Converted {len(images)} pages to images")
            
            # Process images in parallel
            tasks = []
            for idx, image in enumerate(images):
                tasks.append(self._process_page(image, target_language, idx + 1))
            
            # Wait for all pages to be processed
            processed_pages = await asyncio.gather(*tasks)
            logger.info("All pages processed successfully")
            
            # Generate Word document with translated content
            output_path = await self._generate_document(file_path, processed_pages, image_paths)
            logger.info(f"Generated document: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
            
    async def _pdf_to_images(self, pdf_path: str) -> Tuple[List[Image.Image], List[str]]:
        """Convert PDF pages to images with error handling and save to folder."""
        try:
            output_dir = Path('static/extracted_pages')
            output_dir.mkdir(exist_ok=True)
            
            # Convert PDF using PyMuPDF
            logger.info("Converting PDF to images...")
            pdf_document = fitz.open(pdf_path)
            images = []
            image_paths = []
            
            for page_num in range(pdf_document.page_count):
                # Get the page
                page = pdf_document[page_num]
                
                # Convert page to image with high resolution
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                
                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Save the image
                image_path = output_dir / f"page_{page_num + 1:03d}.png"
                img.save(image_path, "PNG", optimize=True)
                logger.info(f"Saved page {page_num + 1} to {image_path}")
                
                images.append(img)
                image_paths.append(str(image_path))
            
            pdf_document.close()
            logger.info(f"Successfully converted PDF to {len(images)} images")
            return images, image_paths
            
        except Exception as e:
            logger.error(f"Error converting PDF to images: {str(e)}")
            raise
            
    async def _process_page(self, image: Image.Image, target_language: str, page_num: int) -> str:
        """Process single page image with GPT-4o for text extraction and translation"""
        try:
            # Save image to bytes
            img_byte_arr = BytesIO()
            image.save(img_byte_arr, format='JPEG')
            img_byte_arr = img_byte_arr.getvalue()
            
            # Encode image to base64
            base64_image = base64.b64encode(img_byte_arr).decode('utf-8')
            
            # Extract and translate text using GPT-4o
            response = await self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": f"""You are a document analysis and translation expert. Your task is to:
                        1. Extract all text content from the image
                        2. Translate the content to {target_language}
                        3. Return a JSON object with both original and translated text
                        4. Maintain paragraph structure and formatting
                        Return format:
                        {{
                            "original": "Original text with paragraphs",
                            "translated": "Translated text with same paragraph structure"
                        }}"""
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": f"Extract and translate the text to {target_language}, maintaining paragraph structure."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                response_format={ "type": "json_object" }
            )
            
            # Parse the response
            result = json.loads(response.choices[0].message.content)
            return result
            
        except Exception as e:
            logger.error(f"Error processing page {page_num}: {str(e)}")
            raise
            
    async def _generate_document(self, original_pdf: str, translated_pages: List[dict], image_paths: List[str]) -> str:
        """Generate Word document with translated content and page images"""
        try:
            # Create output path
            output_path = str(self.temp_dir / f"translated_{os.path.basename(original_pdf)}.docx")
            
            # Create Word document
            doc = Document()
            
            # Set font for the entire document
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Add title
            title = doc.add_heading('Translated Document', 0)
            title.alignment = 1  # Center alignment
            
            # Add each page's content
            for i, (content, image_path) in enumerate(zip(translated_pages, image_paths), 1):
                # Add page header
                doc.add_heading(f'Page {i}', level=1)
                
                # Add the page image
                doc.add_picture(image_path, width=Inches(6.0))
                doc.add_paragraph()  # Add spacing
                
                # Add original text section
                doc.add_heading('Original Text:', level=2)
                doc.add_paragraph(content.get('original', ''))
                doc.add_paragraph()  # Add spacing
                
                # Add translated text section
                doc.add_heading('Translated Text:', level=2)
                doc.add_paragraph(content.get('translated', ''))
                
                # Add page break between pages
                if i < len(translated_pages):
                    doc.add_page_break()
            
            # Save the document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            raise 