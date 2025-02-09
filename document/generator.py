"""Document generation functionality"""

from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from termcolor import colored

class DocumentGenerator:
    """Generate formatted Word documents"""
    
    @staticmethod
    async def create_document(contents: List[str], output_path: str, 
                            output_type: str, target_language: str) -> None:
        """Create a formatted Word document with the content"""
        try:
            doc = Document()
            
            # Add title
            title = f"{'Tweet' if output_type == 'tweet' else 'Transcript'} in {target_language}"
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            for i, content in enumerate(contents, 1):
                # Add section header if multiple contents
                if len(contents) > 1:
                    header = doc.add_paragraph()
                    header_run = header.add_run(f"Content {i}")
                    header_run.font.size = Pt(14)
                    header_run.font.bold = True
                
                # Add the actual content
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(content)
                content_run.font.size = Pt(11)
                
                # Add spacing between sections
                if i < len(contents):
                    doc.add_paragraph()
            
            # Save the document
            doc.save(output_path)
            
        except Exception as e:
            print(colored(f"Error creating document: {str(e)}", "red"))
            raise 